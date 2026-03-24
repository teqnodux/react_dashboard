#!/usr/bin/env python3
"""Quick script to inspect DOCX structure"""
from docx import Document
from pathlib import Path

docx_path = Path(__file__).parent / 'data' / 'dma_docs' / 'input' / 'D001_DMA.docx'
doc = Document(str(docx_path))

print("=" * 80)
print("DOCX STRUCTURE INSPECTION")
print("=" * 80)
print(f"\nTotal paragraphs: {len(doc.paragraphs)}\n")

# Look for section markers and show structure
for i, para in enumerate(doc.paragraphs[:50]):  # First 50 paragraphs
    text = para.text.strip()
    if not text:
        continue

    style = para.style.name
    is_bold = para.runs and para.runs[0].bold if para.runs else False

    # Check for key markers
    is_concise = 'concise' in text.lower() and is_bold
    is_fulsome = 'fulsome' in text.lower() and is_bold
    is_heading = 'Heading' in style

    marker = ""
    if is_concise:
        marker = " <<<< CONCISE MARKER"
    elif is_fulsome:
        marker = " <<<< FULSOME MARKER"
    elif is_heading:
        marker = f" <<<< {style}"

    if marker or i < 10:  # Show first 10 or any with markers
        print(f"{i:3d} | {style:20s} | Bold: {is_bold} | {text[:80]}{marker}")

print("\n" + "=" * 80)
print("Looking for 'Concise Summary' and 'Fulsome Summary' sections...")
print("=" * 80)

concise_found = False
fulsome_found = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'concise summary' in text.lower():
        print(f"✓ Found 'Concise Summary' at paragraph {i}: {text}")
        concise_found = True
    if 'fulsome summary' in text.lower():
        print(f"✓ Found 'Fulsome Summary' at paragraph {i}: {text}")
        fulsome_found = True

if not concise_found:
    print("✗ 'Concise Summary' section NOT found")
if not fulsome_found:
    print("✗ 'Fulsome Summary' section NOT found")

print("\n" + "=" * 80)
